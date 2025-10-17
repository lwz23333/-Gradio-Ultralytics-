import sys
import subprocess
from pathlib import Path

try:
    from docx import Document  # type: ignore
except Exception:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document  # type: ignore


def extract_docx_to_md(path: Path) -> Path:
    doc = Document(str(path))
    lines: list[str] = []

    # paragraphs
    for para in doc.paragraphs:
        text = para.text.replace('\u200b', '').strip()
        if text:
            lines.append(text)

    # tables (flatten as pipe-separated rows)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            if any(cell for cell in row_text):
                lines.append(' | '.join(row_text))

    md_path = path.with_suffix('.md')
    md_path.write_text('\n'.join(lines), encoding='utf-8')
    return md_path


def main() -> None:
    root = Path(__file__).resolve().parent
    files = [
        '模型分析报告.docx',
        '文献分析报告.docx',
        '实验分析报告.docx',
        '技术方案设计报告.docx',
    ]
    for fn in files:
        p = root / fn
        if p.exists():
            out = extract_docx_to_md(p)
            print(f'OK {p.name} -> {out.name}')
        else:
            print(f'MISSING {p.name}')


if __name__ == '__main__':
    main()
